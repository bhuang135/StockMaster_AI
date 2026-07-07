# Filename: stock_data_utils.py
# Updated: RAG-grounded report with direct hyperlinked news and integrated NLP tone
# -------------------------------------------------------------




import os
import re
from io import BytesIO
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Any, Optional

import requests
import yfinance as yf
import pandas as pd
import pandas_ta as ta
import google.generativeai as genai

from rag_chat_pipeline import (
    build_chat_knowledge_base,
    retrieve_chat_documents,
    generate_chat_answer_with_citations
)


from datetime import datetime
from collections import defaultdict
from typing import Any, Dict, List, Optional

import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


def _safe_float(value, default=None):
    """
    Convert int/float/string to float safely.
    Examples:
    '12.5%' -> 12.5
    '1,250.7' -> 1250.7
    None -> default
    """
    if value is None:
        return default

    if isinstance(value, (int, float)):
        return float(value)

    if isinstance(value, str):
        cleaned = value.replace(",", "").replace("%", "").strip()
        match = re.search(r"-?\d+(\.\d+)?", cleaned)
        if match:
            try:
                return float(match.group())
            except Exception:
                return default

    return default


def _parse_market_cap_million(market_cap_value) -> Optional[float]:
    """
    Convert Market Cap text to 'million USD' float.
    Example:
    '250,000 Million USD' -> 250000.0
    """
    if market_cap_value is None:
        return None

    if isinstance(market_cap_value, (int, float)):
        return float(market_cap_value) / 1e6

    if isinstance(market_cap_value, str):
        num = _safe_float(market_cap_value, default=None)
        if num is not None:
            return num

    return None

# ----------------------------
# 0) Gemini configuration
# ----------------------------
# NOTE: We intentionally do NOT hardcode a Gemini API Key in this file.
# The key is supplied by the user at runtime (typed into the ticker search
# UI) and passed into these functions as the `api_key` argument. It is only
# cached in the browser's session storage on the frontend and is never
# written into this codebase. GEMINI_API_KEY env var is kept only as an
# optional fallback for local/dev use.
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip()

def _get_gemini_model(api_key: Optional[str] = None):
    """
    Build a Gemini model client using a user-supplied API key.
    Falls back to the GEMINI_API_KEY environment variable only if the
    caller didn't provide one (useful for local dev), but the normal
    path is for `api_key` to come from what the user typed in the UI.
    """
    key_to_use = (api_key or GEMINI_API_KEY or "").strip()
    if not key_to_use:
        return None
    genai.configure(api_key=key_to_use)
    return genai.GenerativeModel("gemini-2.5-flash") 

# ----------------------------
# 1) Market / fundamentals
# ----------------------------
def get_stock_data(ticker: str) -> Tuple[pd.DataFrame, Dict[str, Any]]:
    try:
        df = yf.download(
            ticker,
            period="5y",
            interval="1d",
            progress=False,
            auto_adjust=True
        )

        if df is None or df.empty:
            return pd.DataFrame(), {}

        if isinstance(df.columns, pd.MultiIndex):
            df.columns = df.columns.get_level_values(0)

        df = df.dropna()
        df.reset_index(inplace=True)
        df.set_index("Date", inplace=True)
        df.columns = [str(col) for col in df.columns]

        df.ta.rsi(append=True)

        stock = yf.Ticker(ticker)
        info = stock.info if hasattr(stock, "info") else {}

        def _format_market_cap(val):
            try:
                if val:
                    val_m = float(val) / 1e6
                    return f"{val_m:,.0f} Million USD"
            except Exception:
                pass
            return "N/A"

        def _safe_join_parts(*parts):
            cleaned = [str(p).strip() for p in parts if p not in [None, "", "N/A"]]
            return ", ".join(cleaned) if cleaned else "N/A"

        def _extract_ceo_name(company_officers):
            try:
                if not company_officers:
                    return "N/A"

                # ���� title �t CEO / Chief Executive Officer
                for officer in company_officers:
                    title = str(officer.get("title", "")).lower()
                    name = officer.get("name")
                    if ("chief executive officer" in title or title == "ceo" or " ceo" in title or "ceo " in title):
                        return name or "N/A"

                # ����Gtitle �̦� president and ceo ����
                for officer in company_officers:
                    title = str(officer.get("title", "")).lower()
                    name = officer.get("name")
                    if "ceo" in title:
                        return name or "N/A"

                return "N/A"
            except Exception:
                return "N/A"

        def _extract_cfo_name(company_officers):
            try:
                if not company_officers:
                    return "N/A"

                for officer in company_officers:
                    title = str(officer.get("title", "")).lower()
                    name = officer.get("name")
                    if ("chief financial officer" in title or title == "cfo" or " cfo" in title or "cfo " in title):
                        return name or "N/A"

                for officer in company_officers:
                    title = str(officer.get("title", "")).lower()
                    name = officer.get("name")
                    if "cfo" in title:
                        return name or "N/A"

                return "N/A"
            except Exception:
                return "N/A"

        def _extract_founder(info_dict):
            # Yahoo Finance ���@�wí�w���� founder�A�o�̫O�u�B�z
            for key in ["founder", "Founders", "founders"]:
                if key in info_dict and info_dict.get(key):
                    return info_dict.get(key)
            return "N/A"

        company_officers = info.get("companyOfficers", []) or []

        fundamentals = {
            "Symbol": ticker.upper(),
            "Company Name": info.get("longName", ticker.upper()),
            "CEO": _extract_ceo_name(company_officers),
            "CFO": _extract_cfo_name(company_officers),
            "Founder": _extract_founder(info),
            "P/E Ratio (TTM)": info.get("trailingPE"),
            "Revenue Growth (YoY)": info.get("revenueGrowth"),
            "Profit Margin": info.get("profitMargins"),
            "ROE (Return on Equity)": info.get("returnOnEquity"),
            "Debt/Equity": info.get("debtToEquity"),
            "Market Cap": _format_market_cap(info.get("marketCap")),
            "Sector": info.get("sector"),
            "Industry": info.get("industry"),
            "Business Summary": info.get("longBusinessSummary"),
            "Website": info.get("website"),
            "Headquarters": _safe_join_parts(
                info.get("city"),
                info.get("state"),
                info.get("zip")
            ),
            "Country": info.get("country", "N/A"),
            "Full Time Employees": info.get("fullTimeEmployees", "N/A"),
            "Exchange": info.get("exchange", "N/A"),
            "Currency": info.get("currency", "N/A"),
            "Data Source": "Yahoo Finance (yfinance)",
            "Data Period": f"5y (Until {datetime.now().strftime('%Y-%m-%d')})"
        }

        for key in ["Revenue Growth (YoY)", "Profit Margin", "ROE (Return on Equity)"]:
            val = fundamentals.get(key)
            if isinstance(val, (int, float)):
                fundamentals[key] = f"{val * 100:.2f}%"

        return df, fundamentals

    except Exception:
        return pd.DataFrame(), {}

def get_financial_statements(ticker: str) -> pd.DataFrame:
    try:
        stock = yf.Ticker(ticker)
        df_financials = stock.financials
        if df_financials is None or df_financials.empty:
            return pd.DataFrame()

        metrics = ["Total Revenue", "Gross Profit", "Operating Income", "Net Income"]
        available_metrics = [m for m in metrics if m in df_financials.index]
        if not available_metrics:
            return pd.DataFrame()

        df_pl = df_financials.loc[available_metrics].transpose()
        df_pl.index = pd.to_datetime(df_pl.index).year
        df_pl = df_pl.sort_index(ascending=True).apply(pd.to_numeric, errors="coerce") / 1e6

        df_res = df_pl.reset_index().rename(columns={"index": "Year"})
        df_res["Year"] = df_res["Year"].astype(int)
        return df_res
    except Exception as e:
        print(f"Financial data retrieval failed: {e}")
        return pd.DataFrame()


# ----------------------------
# 1b) Gemini fallback for missing/N-A fundamentals
# ----------------------------
# Maps each fundamentals field that can legitimately come back missing from
# Yahoo Finance to a short human-readable description used when asking
# Gemini to look it up. Symbol / Company Name / Data Source / Data Period
# are deliberately excluded — they're either always set programmatically or
# aren't real facts to search for.
_FUNDAMENTALS_GEMINI_LOOKUP_LABELS: Dict[str, str] = {
    "CEO": "Chief Executive Officer's name",
    "CFO": "Chief Financial Officer's name",
    "Founder": "founder's name",
    "Sector": "market sector",
    "Industry": "industry",
    "Headquarters": "headquarters city/state",
    "Country": "country of headquarters",
    "Website": "official website URL",
    "Full Time Employees": "approximate number of full-time employees",
    "Exchange": "primary stock exchange",
    "Currency": "reporting currency",
    "Business Summary": "one to two sentence description of what the company does",
    "P/E Ratio (TTM)": "trailing twelve months P/E ratio",
    "Revenue Growth (YoY)": "most recent year-over-year revenue growth percentage",
    "Profit Margin": "most recent net profit margin percentage",
    "ROE (Return on Equity)": "most recent return on equity percentage",
    "Debt/Equity": "most recent debt-to-equity ratio",
    "Market Cap": "current market capitalization",
}

_MISSING_VALUE_MARKERS = {"", "n/a", "na", "none", "unknown", "-", "null"}


def _is_missing_value(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return value.strip().lower() in _MISSING_VALUE_MARKERS
    return False


def fill_missing_fundamentals_with_gemini(
    fundamentals: Dict[str, Any],
    ticker: str,
    api_key: Optional[str] = None,
) -> Dict[str, Any]:
    """
    For any fundamentals field that came back missing / "N/A" from Yahoo
    Finance, try to fill it in with a single Gemini + Google Search
    grounded lookup (one API call covers every missing field at once, to
    keep this cheap).

    Safety note (same trust boundary as the news sources): the whole batch
    is discarded unless the response actually came back grounded in real
    Google Search results. We do not accept an ungrounded guess for a
    company fact like "CEO name" — better to keep showing "N/A" than to
    risk showing a hallucinated fact. Any field that *is* filled in gets
    recorded in `fundamentals["_gemini_filled_fields"]` so the UI can make
    clear which values came from a live Gemini lookup rather than the
    primary Yahoo Finance data source.
    """
    key_to_use = (api_key or GEMINI_API_KEY or "").strip()
    if not key_to_use or not fundamentals:
        return fundamentals

    missing_fields = [
        field for field in _FUNDAMENTALS_GEMINI_LOOKUP_LABELS
        if field in fundamentals and _is_missing_value(fundamentals.get(field))
    ]
    if not missing_fields:
        return fundamentals

    company_name = fundamentals.get("Company Name") or ticker.upper()
    if _is_missing_value(company_name):
        company_name = ticker.upper()

    try:
        genai.configure(api_key=key_to_use)
        model = genai.GenerativeModel("gemini-2.5-flash")

        field_list_md = "\n".join(
            f"- {field}: {_FUNDAMENTALS_GEMINI_LOOKUP_LABELS[field]}"
            for field in missing_fields
        )

        prompt = f"""
Search for the following missing data points about {company_name} ({ticker.upper()}).

{field_list_md}

Reply with ONLY one line per field, in exactly this format, with no extra
commentary before or after:
Field Name: value

If you cannot verify a field with a real, current source, reply with
"Field Name: UNKNOWN" for that field instead of guessing.
"""

        try:
            response = model.generate_content(prompt, tools="google_search")
        except Exception as tool_err:
            print(f"Gemini grounding tool unavailable for fundamentals fill: {tool_err}")
            return fundamentals

        candidates = getattr(response, "candidates", None) or []
        if not candidates:
            return fundamentals

        grounding_metadata = getattr(candidates[0], "grounding_metadata", None)
        if not grounding_metadata:
            # Not actually grounded in Google Search -> don't trust any of it.
            return fundamentals

        response_text = getattr(response, "text", "") or ""
        filled_fields: List[str] = []

        for line in response_text.splitlines():
            if ":" not in line:
                continue
            field_name, _, value = line.partition(":")
            field_name = field_name.strip().lstrip("-").strip()
            value = value.strip()
            if (
                field_name in missing_fields
                and value
                and value.upper() != "UNKNOWN"
            ):
                fundamentals[field_name] = value
                filled_fields.append(field_name)

        if filled_fields:
            fundamentals["_gemini_filled_fields"] = filled_fields

        return fundamentals

    except Exception as e:
        print(f"Gemini fundamentals fill failed: {e}")
        return fundamentals


def dynamic_valuation_model(
    pe_ratio,
    revenue_growth,
    profit_margin,
    market_cap,
    sector=None,
    industry=None
):
    """
    Return:
        model_config = {
            "model_desc": str,
            "base_upside": float,
            "base_downside": float,
            "sentiment_sensitivity": float,
            "risk_penalty": float
        }

    Interpretation:
    - base_upside: neutral-to-positive upside capacity
    - base_downside: neutral-to-negative downside capacity
    - sentiment_sensitivity: how strongly sentiment affects price scenarios
    - risk_penalty: structural discount for weak profitability / risky profile
    """

    model_config = {
        "model_desc": "Intermediate Growth Model",
        "base_upside": 0.10,
        "base_downside": 0.08,
        "sentiment_sensitivity": 0.60,
        "risk_penalty": 0.00
    }

    try:
        # ----------------------------
        # 1) Parse inputs
        # ----------------------------
        pe = _safe_float(pe_ratio, default=20.0)
        rev_growth_pct = _safe_float(revenue_growth, default=10.0)   # e.g. "12.50%" -> 12.5
        margin_pct = _safe_float(profit_margin, default=12.0)        # e.g. "25.00%" -> 25.0
        market_cap_m = _parse_market_cap_million(market_cap)         # in million USD

        rev_growth = rev_growth_pct / 100.0
        profit_margin_ratio = margin_pct / 100.0

        sector = (sector or "").strip()
        industry = (industry or "").strip()

        # ----------------------------
        # 2) Base company profile
        # ----------------------------
        # Inference:
        # high PE + high growth => growth stock
        # low PE or low growth => mature / value stock
        if pe > 30 and rev_growth > 0.15:
            model_config.update({
                "model_desc": "High Growth Model",
                "base_upside": 0.18,
                "base_downside": 0.12,
                "sentiment_sensitivity": 0.90
            })
        elif pe < 20 or rev_growth < 0.05:
            model_config.update({
                "model_desc": "Mature / Value Model",
                "base_upside": 0.08,
                "base_downside": 0.06,
                "sentiment_sensitivity": 0.40
            })
        else:
            model_config.update({
                "model_desc": "Intermediate Growth Model",
                "base_upside": 0.10,
                "base_downside": 0.08,
                "sentiment_sensitivity": 0.60
            })

        # ----------------------------
        # 3) Sector / industry overlay
        # ----------------------------
        if sector in ["Technology", "Semiconductors"] or "Semiconductor" in industry:
            model_config.update({
                "model_desc": "AI / Tech Growth Model",
                "base_upside": max(model_config["base_upside"], 0.20),
                "base_downside": max(model_config["base_downside"], 0.14),
                "sentiment_sensitivity": 1.10
            })

        elif sector in ["Financial Services"]:
            model_config.update({
                "model_desc": "Bank / Financial Cycle Model",
                "base_upside": 0.09,
                "base_downside": 0.09,
                "sentiment_sensitivity": 0.45
            })

        elif sector in ["Energy", "Basic Materials"]:
            model_config.update({
                "model_desc": "Commodity Cycle Model",
                "base_upside": 0.14,
                "base_downside": 0.13,
                "sentiment_sensitivity": 0.70
            })

        elif sector in ["Utilities", "Consumer Defensive", "Healthcare"]:
            model_config.update({
                "model_desc": "Defensive Stability Model",
                "base_upside": 0.07,
                "base_downside": 0.05,
                "sentiment_sensitivity": 0.30
            })

        # ----------------------------
        # 4) Market cap overlay
        # market_cap_m is in million USD
        # 200B = 200,000 million
        # 10B  = 10,000 million
        # ----------------------------
        if market_cap_m is not None:
            if market_cap_m > 200_000:
                # mega-cap: more stable, less upside/downside swing
                model_config["model_desc"] += " | Mega-Cap Stability Overlay"
                model_config["base_upside"] *= 0.90
                model_config["base_downside"] *= 0.85
                model_config["sentiment_sensitivity"] *= 0.85

            elif market_cap_m < 10_000:
                # small-cap: more volatile, sentiment hits harder
                model_config["model_desc"] += " | Small-Cap Volatility Overlay"
                model_config["base_upside"] *= 1.20
                model_config["base_downside"] *= 1.25
                model_config["sentiment_sensitivity"] *= 1.20

        # ----------------------------
        # 5) Profitability overlay
        # ----------------------------
        if profit_margin_ratio < 0:
            model_config["model_desc"] += " | Loss-Making Risk Overlay"
            model_config["risk_penalty"] += 0.05
            model_config["base_downside"] += 0.04

        elif profit_margin_ratio > 0.20:
            model_config["model_desc"] += " | High-Margin Quality Overlay"
            model_config["base_upside"] += 0.02

        return model_config

    except Exception:
        return model_config

def simulated_valuation(df: pd.DataFrame, fundamentals: Dict[str, Any], api_key: Optional[str] = None):
    """
    Return:
        current_close, positive_price, neutral_price, negative_price, model_desc, valuation_debug
    """
    if df.empty:
        return 0, 0, 0, 0, "N/A", {}

    current_close = float(df["Close"].iloc[-1])

    # ----------------------------
    # 1) Collect sentiment
    # ----------------------------
    sentiment_total = 0.0
    sentiment_summary = {}
    try:
        ticker = fundamentals.get("Symbol")
        news_items = get_recent_news(ticker, 10, api_key=api_key)
        sentiment_summary = summarize_news_sentiment(news_items)
        sentiment_total = sentiment_summary["clamped_sentiment_sum"]
    except Exception:
        news_items = []
        sentiment_summary = summarize_news_sentiment([])

    # ----------------------------
    # 2) Select model config
    # ----------------------------
    model_config = dynamic_valuation_model(
        pe_ratio=fundamentals.get("P/E Ratio (TTM)"),
        revenue_growth=fundamentals.get("Revenue Growth (YoY)"),
        profit_margin=fundamentals.get("Profit Margin"),
        market_cap=fundamentals.get("Market Cap"),
        sector=fundamentals.get("Sector"),
        industry=fundamentals.get("Industry")
    )

    base_upside = model_config["base_upside"]
    base_downside = model_config["base_downside"]
    sentiment_sensitivity = model_config["sentiment_sensitivity"]
    risk_penalty = model_config["risk_penalty"]
    model_desc = model_config["model_desc"]

    # ----------------------------
    # 3) Additional risk overlays
    # ----------------------------
    negative_impact = risk_penalty

    try:
        debt = _safe_float(fundamentals.get("Debt/Equity"), default=None)
        if debt is not None and debt > 150:
            negative_impact += 0.04
    except Exception:
        pass

    try:
        rsi_cols = [c for c in df.columns if "RSI" in c]
        if rsi_cols:
            latest_rsi = float(df[rsi_cols[0]].iloc[-1])
            if latest_rsi > 75:
                negative_impact += 0.03
            elif latest_rsi < 30:
                # oversold rebound benefit
                base_upside += 0.02
    except Exception:
        pass

    # prevent excessive penalty
    negative_impact = min(negative_impact, 0.20)

    # ----------------------------
    # 4) Scenario construction
    # ----------------------------
    # Assumption:
    # - Neutral = structural model only
    # - Positive = structural model + positive sentiment boost
    # - Negative = structural model - downside - risk
    #
    # sentiment_total can be negative or positive
    positive_sentiment_boost = max(sentiment_total, 0.0) * sentiment_sensitivity
    negative_sentiment_boost = abs(min(sentiment_total, 0.0)) * sentiment_sensitivity

    neutral_return = max((base_upside * 0.50) - (negative_impact * 0.50), -0.10)
    positive_return = max(base_upside + positive_sentiment_boost - negative_impact * 0.30, -0.20)
    negative_return = min(-(base_downside + negative_sentiment_boost + negative_impact), 0.0)

    neutral_price = current_close * (1 + neutral_return)
    positive_price = current_close * (1 + positive_return)
    negative_price = current_close * (1 + negative_return)


    valuation_debug = {
        "model_desc": model_desc,
        "base_upside": base_upside,
        "base_downside": base_downside,
        "sentiment_sensitivity": sentiment_sensitivity,
        "risk_penalty": risk_penalty,
        "negative_impact": negative_impact,
        "positive_sentiment_boost": positive_sentiment_boost,
        "negative_sentiment_boost": negative_sentiment_boost,
        "neutral_return": neutral_return,
        "positive_return": positive_return,
        "negative_return": negative_return,
        "sentiment_summary": sentiment_summary
        }

    return (
        current_close,
        positive_price,
        neutral_price,
        negative_price,
        model_desc,
        valuation_debug
        )

# ----------------------------
# 3) News & Options Retrieval
# ----------------------------
#
# --- Sentiment lexicon -------------------------------------------------
# NOTE ON THE FIX (2026-07):
# The previous implementation tokenized text into single alphabetic words
# (regex `[a-zA-Z]+`) and then checked *set membership* against a lexicon
# that is mostly made of hyphenated MULTI-WORD phrases such as
# "record-high", "revenue-growth", "earnings-miss", "guidance-cut", etc.
# Tokenizing on `[a-zA-Z]+` strips the hyphen, turning "record-high" into
# two separate tokens "record" and "high" — neither of which is itself in
# the lexicon (only the hyphenated compound is). As a result, ~80% of the
# lexicon could never match real headlines, and most articles silently
# scored as "Neutral" regardless of their actual content, which broke the
# whole "news sentiment -> positive/negative price scenario" pipeline.
#
# THE FIX:
# Instead of tokenizing and checking set membership, each lexicon term is
# compiled into a regex pattern where its hyphen becomes a flexible
# separator (`[\s-]+`), so "record-high" matches BOTH "record-high" and
# "record high" (how headlines actually write it), and the pattern is
# searched directly against the (lightly normalized) article text with
# word boundaries. A lightweight negation window is also added so phrases
# like "not profitable" or "no losses" flip the polarity of the term they
# precede, which the previous version could not do at all.
POSITIVE_SENTIMENT_TERMS = {
    "advance", "advancing", "advancement",
    "boom", "booming",
    "strengthen", "strengthening",
    "stabilize", "stabilized",
    "improve", "improved", "improvement",
    "enhance", "enhanced", "enhancement",
    "achieve", "achievement",
    "deliver", "delivered",
    "solid", "robust", "resilient",
    "stronger", "uptrend",
    "milestone", "breakthrough",
    "record-high", "all-time-high",

    "expedite", "accelerated",
    "scale", "scaled", "scaling",
    "diversify", "diversified",
    "optimize", "optimized",
    "modernize", "modernized",

    "profitability", "profitable",
    "earnings-beat", "revenue-growth",
    "margin-expansion",
    "cash-flow-positive",

    "approval-granted",
    "authorization",
    "greenlight",
    "certification",
    "compliance-approved",

    "collaboration",
    "strategic-alliance",
    "joint-venture",
    "merger-successful",
    "integration-successful",

    "investment", "funding",
    "capital-injection",
    "venture-backed",
    "ipo-successful",
    "oversubscribed",

    "upgrade-rating",
    "price-target-raised",
    "analyst-confidence",

    "demand-surges",
    "sales-growth",
    "customer-growth",
    "user-growth",

    "market-share-gain",
    "leadership-position",

    "cost-reduction",
    "efficiency-gain",

    "product-launch-success",
    "innovation-driven",
    "technology-breakthrough",

    "expansion-plan",
    "global-expansion",
    "market-entry",

    "recovery-phase",
    "turnaround-success",

    "restructure-success",
    "debt-reduction",

    "share-buyback",
    "dividend-increase",

    "positive-outlook",
    "guidance-raised",

    "outlook-improved",
    "forecast-upgrade",

    "strategic-win",
    "contract-awarded",

    "talent-acquisition",

    "award-winning",

    "strong-demand",

    "healthy-balance-sheet",

    "liquidity-improved",

    "valuation-upside",

    "bull-run",

    "market-confidence",

    "upside-potential",
}

NEGATIVE_SENTIMENT_TERMS = {
    "deteriorate", "deterioration",
    "weaken", "weakening",
    "erosion",

    "pressure", "downward-pressure",

    "underestimate",
    "shortfall",
    "earnings-miss",

    "loss", "losses",

    "margin-compression",

    "cash-burn",

    "liquidity-crunch",

    "default-risk",

    "credit-risk",

    "write-down",
    "write-off",

    "impairment",

    "volatility",

    "instability",

    "disruption",

    "shutdown",

    "halted",

    "suspended",

    "terminated",

    "delay", "delayed",

    "postponed",

    "canceled",

    "recall-issue",

    "defect",

    "safety-concern",

    "compliance-failure",

    "violation",

    "non-compliance",

    "fraud",

    "scandal",

    "misconduct",

    "allegation",

    "investigated",

    "charged",

    "lawsuit-filed",

    "settlement",

    "legal-risk",

    "regulatory-risk",

    "antitrust",

    "sanction",

    "blacklist",

    "restriction",

    "ban",

    "export-ban",

    "trade-war",

    "tariff-impact",

    "supply-chain-disruption",

    "shortage",

    "bottleneck",

    "inventory-glut",

    "oversupply",

    "slow-sales",

    "demand-drop",

    "market-share-loss",

    "competition-intensifies",

    "price-war",

    "cost-increase",

    "inflation-impact",

    "expense-surge",

    "layoffs-announced",

    "job-cuts",

    "workforce-reduction",

    "restructuring-charge",

    "plant-closure",

    "bankruptcy-filing",

    "insolvency",

    "delisting-risk",

    "downgraded-rating",

    "negative-outlook",

    "guidance-cut",

    "forecast-lowered",

    "profit-warning",

    "revenue-decline",

    "sales-slump",

    "earnings-drop",

    "bear-market",

    "selloff",

    "panic-selling",

    "market-turmoil",

    "risk-off",

    "uncertain-outlook",

    "geopolitical-risk",

    "cyberattack",

    "data-breach",

    "system-failure",

    "outage",

    "security-flaw",

    "customer-loss",

    "brand-damage",

    "reputation-risk",

    "contract-loss",
}

# Words that flip the polarity of a sentiment term that follows shortly
# after them, e.g. "not profitable", "no losses", "failed to deliver".
_NEGATION_CUE_WORDS = {
    "not", "no", "never", "without", "fails", "fail", "failed", "failing",
    "unable", "cannot", "lacks", "lacking", "denies", "denied", "avoid", "avoids",
}
_NEGATION_WINDOW_WORDS = 4  # how many words before a match to scan for a negation cue


def _compile_term_patterns(terms: set) -> List[Tuple[str, "re.Pattern"]]:
    """
    Compile each lexicon term into a regex that matches the term whether it
    appears hyphenated ("record-high") or as separate words
    ("record high") — this is what the old tokenizer-based matcher could
    not do for any multi-word term.
    """
    compiled = []
    for term in terms:
        parts = [re.escape(p) for p in term.split("-")]
        pattern_str = r"\b" + r"[\s-]+".join(parts) + r"\b"
        compiled.append((term, re.compile(pattern_str, re.IGNORECASE)))
    return compiled


_POSITIVE_TERM_PATTERNS = _compile_term_patterns(POSITIVE_SENTIMENT_TERMS)
_NEGATIVE_TERM_PATTERNS = _compile_term_patterns(NEGATIVE_SENTIMENT_TERMS)


def _has_negation_before(text_lower: str, match_start: int) -> bool:
    """Check whether a negation cue word appears shortly before a match."""
    preceding_text = text_lower[:match_start]
    preceding_words = re.findall(r"[a-z']+", preceding_text)[-_NEGATION_WINDOW_WORDS:]
    return any(w in _NEGATION_CUE_WORDS for w in preceding_words)


def _get_local_sentiment_tag(text: str) -> Dict[str, Any]:
    """
    Score a piece of text (headline and/or description) using the local
    keyword lexicon. Returns a raw integer score, a capped per-article
    weight, and the list of matched terms (useful for debugging /
    transparency in the generated report).
    """
    text_lower = (text or "").lower()

    matched_pos: List[str] = []
    matched_neg: List[str] = []
    raw_score = 0

    for term, pattern in _POSITIVE_TERM_PATTERNS:
        for m in pattern.finditer(text_lower):
            if _has_negation_before(text_lower, m.start()):
                raw_score -= 1
                matched_neg.append(f"{term} (negated)")
            else:
                raw_score += 1
                matched_pos.append(term)

    for term, pattern in _NEGATIVE_TERM_PATTERNS:
        for m in pattern.finditer(text_lower):
            if _has_negation_before(text_lower, m.start()):
                raw_score += 1
                matched_pos.append(f"{term} (negated)")
            else:
                raw_score -= 1
                matched_neg.append(term)

    # normalize into a capped score so one article does not dominate
    # inference: capped between -0.03 and +0.03 to stay consistent with the
    # existing valuation framework (see summarize_news_sentiment / clamp).
    if raw_score > 0:
        tone = "Positive"
        sentiment_weight = min(raw_score * 0.01, 0.03)
    elif raw_score < 0:
        tone = "Negative"
        sentiment_weight = max(raw_score * 0.01, -0.03)
    else:
        tone = "Neutral"
        sentiment_weight = 0.0

    return {
        "tone_tag": tone,
        "raw_score": raw_score,
        "sentiment_weight": sentiment_weight,
        "matched_positive_terms": matched_pos,
        "matched_negative_terms": matched_neg,
    }



def _get_yahoo_finance_news(ticker: str, max_items: int = 10) -> List[Dict[str, Any]]:
    """
    Source 1: Yahoo Finance, via the yfinance library that's already a
    dependency of this app. No extra API key needed, so this source is
    always available as a baseline even if the Gemini key hasn't been
    entered yet.

    yfinance's `Ticker.news` schema has changed across versions — newer
    versions nest the real fields under an item's "content" key, older
    versions put them at the top level. Both shapes are handled here.
    """
    items: List[Dict[str, Any]] = []

    try:
        raw_news = yf.Ticker(ticker).news or []
    except Exception as e:
        print(f"Yahoo Finance source failed: {e}")
        return items

    for a in raw_news[:max_items]:
        try:
            content = a.get("content") if isinstance(a.get("content"), dict) else a

            title = content.get("title", "") or ""
            summary = content.get("summary", "") or content.get("description", "") or ""

            publisher = "Yahoo Finance"
            provider = content.get("provider")
            if isinstance(provider, dict) and provider.get("displayName"):
                publisher = provider.get("displayName")
            elif isinstance(a.get("publisher"), str) and a.get("publisher"):
                publisher = a.get("publisher")

            link = ""
            click_through = content.get("clickThroughUrl") or content.get("canonicalUrl")
            if isinstance(click_through, dict):
                link = click_through.get("url", "") or ""
            if not link:
                link = a.get("link", "") or ""

            dt = None
            date_str = "N/A"
            pub_date_raw = content.get("pubDate") or a.get("providerPublishTime")
            if isinstance(pub_date_raw, str):
                try:
                    dt = datetime.fromisoformat(pub_date_raw.replace("Z", ""))
                    date_str = dt.strftime("%Y-%m-%d")
                except Exception:
                    date_str = pub_date_raw[:10] if len(pub_date_raw) >= 10 else "N/A"
            elif isinstance(pub_date_raw, (int, float)):
                try:
                    dt = datetime.utcfromtimestamp(pub_date_raw)
                    date_str = dt.strftime("%Y-%m-%d")
                except Exception:
                    pass

            if not title:
                continue

            sentiment_source_text = f"{title}. {summary}".strip()
            sentiment_result = _get_local_sentiment_tag(sentiment_source_text)

            items.append({
                "title": title,
                "publisher": publisher,
                "date": date_str,
                "published_datetime": dt,
                "link": link or "#",
                "tone_tag": sentiment_result["tone_tag"],
                "raw_score": sentiment_result["raw_score"],
                "sentiment_weight": sentiment_result["sentiment_weight"],
                "matched_positive_terms": sentiment_result["matched_positive_terms"],
                "matched_negative_terms": sentiment_result["matched_negative_terms"],
                "source_type": "yahoo_finance",
            })
        except Exception:
            # skip malformed individual items rather than failing the source
            continue

    return items


def _get_gemini_grounded_news(
    ticker: str,
    company_name: str,
    api_key: Optional[str] = None,
    max_items: int = 6
) -> List[Dict[str, Any]]:
    """
    Source 2: Gemini + Google Search grounding, using the user's own Gemini
    API key. This queries Gemini with the `google_search` grounding tool so
    it can surface current news beyond both the local lexicon's training
    data and Yahoo Finance's coverage.

    Safety note: we ONLY keep results that come back with real grounding
    citations (an actual source URL from Gemini's grounding_metadata). If
    the SDK/model combo doesn't support grounding, or the response isn't
    actually grounded, we return an empty list instead of trusting
    Gemini's free-text — this avoids treating a possible hallucination as
    "news".
    """
    items: List[Dict[str, Any]] = []

    key_to_use = (api_key or GEMINI_API_KEY or "").strip()
    if not key_to_use:
        return items

    try:
        genai.configure(api_key=key_to_use)
        model = genai.GenerativeModel("gemini-2.5-flash")

        prompt = (
            f"Search for the latest notable stock market news from the past 14 days "
            f"about {company_name} ({ticker}). "
            f"List up to {max_items} distinct, concrete news items as short factual "
            f"one-sentence summaries (earnings, guidance, product launches, legal/"
            f"regulatory events, executive changes, analyst actions, etc.)."
        )

        try:
            response = model.generate_content(prompt, tools="google_search")
        except Exception as tool_err:
            # Tool name/support can differ across SDK/model versions. If the
            # grounding tool itself isn't accepted, don't fall back to an
            # ungrounded call — we can't verify an ungrounded answer as news.
            print(f"Gemini grounding tool unavailable: {tool_err}")
            return items

        candidates = getattr(response, "candidates", None) or []
        if not candidates:
            return items

        grounding_metadata = getattr(candidates[0], "grounding_metadata", None)
        if not grounding_metadata:
            # Model answered without actually grounding in Google Search —
            # treat as unverifiable and skip rather than risk fabricated news.
            return items

        grounding_chunks = getattr(grounding_metadata, "grounding_chunks", None) or []
        today_str = datetime.utcnow().strftime("%Y-%m-%d")

        for chunk in grounding_chunks[:max_items]:
            web = getattr(chunk, "web", None)
            if not web:
                continue

            uri = getattr(web, "uri", "") or ""
            if not uri:
                continue

            title = getattr(web, "title", "") or f"{company_name} ({ticker}) — Gemini search result"

            sentiment_result = _get_local_sentiment_tag(title)

            items.append({
                "title": title,
                "publisher": "Gemini (Google Search grounding)",
                "date": today_str,
                "published_datetime": None,
                "link": uri,
                "tone_tag": sentiment_result["tone_tag"],
                "raw_score": sentiment_result["raw_score"],
                "sentiment_weight": sentiment_result["sentiment_weight"],
                "matched_positive_terms": sentiment_result["matched_positive_terms"],
                "matched_negative_terms": sentiment_result["matched_negative_terms"],
                "source_type": "gemini_grounding",
            })

        return items

    except Exception as e:
        print(f"Gemini grounded news source failed: {e}")
        return items


def get_recent_news(ticker: str, max_items: int = 10, api_key: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Aggregate recent news from TWO independent sources for broader, more
    resilient coverage:

      1) Yahoo Finance — via yfinance, always available, no extra key
      2) Gemini         — Google Search-grounded query using the user's own
                           Gemini API key (typed into the search UI)

    Each source is independently best-effort: if one is unavailable or
    fails, the other still contributes, so the pipeline degrades
    gracefully instead of returning nothing. Results are de-duplicated by
    normalized title (both sources sometimes cover the same story) and
    sorted by recency before being truncated to `max_items`.
    """
    company_name = ticker.upper()
    try:
        _, fundamentals = get_stock_data(ticker)
        if fundamentals:
            company_name = fundamentals.get("Company Name", ticker.upper()) or ticker.upper()
    except Exception:
        pass

    all_items: List[Dict[str, Any]] = []
    all_items.extend(_get_yahoo_finance_news(ticker, max_items))
    all_items.extend(_get_gemini_grounded_news(ticker, company_name, api_key=api_key, max_items=max_items))

    # de-duplicate by normalized title (different sources often carry the
    # same story with slightly different punctuation/casing)
    seen_titles = set()
    deduped: List[Dict[str, Any]] = []
    for item in all_items:
        norm_title = re.sub(r"[^a-z0-9]+", "", (item.get("title") or "").lower())
        if not norm_title or norm_title in seen_titles:
            continue
        seen_titles.add(norm_title)
        deduped.append(item)

    # newest first; items with an unknown timestamp sort last
    deduped.sort(
        key=lambda it: it.get("published_datetime") or datetime.min,
        reverse=True
    )

    return deduped[:max_items]


def summarize_news_sentiment(news_items: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Aggregate sentiment diagnostics for reporting.

    Returns:
        {
            "article_count": int,
            "positive_count": int,
            "neutral_count": int,
            "negative_count": int,
            "raw_sentiment_sum": float,
            "clamped_sentiment_sum": float,
            "average_sentiment_weight": float,
            "sentiment_label": str,
            "sentiment_explanation": str,
            "methodology_text": str
        }
    """
    if not news_items:
        return {
            "article_count": 0,
            "positive_count": 0,
            "neutral_count": 0,
            "negative_count": 0,
            "raw_sentiment_sum": 0.0,
            "clamped_sentiment_sum": 0.0,
            "average_sentiment_weight": 0.0,
            "sentiment_label": "Neutral",
            "sentiment_explanation": "No recent news was retrieved, so sentiment impact was treated as neutral.",
            "methodology_text": (
                "Headline/description-based local sentiment scoring was unavailable because no news articles were retrieved."
            )
        }

    positive_count = sum(1 for n in news_items if n.get("tone_tag") == "Positive")
    neutral_count  = sum(1 for n in news_items if n.get("tone_tag") == "Neutral")
    negative_count = sum(1 for n in news_items if n.get("tone_tag") == "Negative")

    raw_sentiment_sum = sum(float(n.get("sentiment_weight", 0.0)) for n in news_items)
    clamped_sentiment_sum = max(min(raw_sentiment_sum, 0.15), -0.15)
    average_sentiment_weight = raw_sentiment_sum / len(news_items) if news_items else 0.0

    if clamped_sentiment_sum >= 0.06:
        sentiment_label = "Moderately Positive"
    elif clamped_sentiment_sum > 0:
        sentiment_label = "Slightly Positive"
    elif clamped_sentiment_sum <= -0.06:
        sentiment_label = "Moderately Negative"
    elif clamped_sentiment_sum < 0:
        sentiment_label = "Slightly Negative"
    else:
        sentiment_label = "Neutral"

    sentiment_explanation = (
        f"Out of {len(news_items)} news items, {positive_count} were positive, "
        f"{neutral_count} were neutral, and {negative_count} were negative. "
        f"The raw headline sentiment sum was {raw_sentiment_sum:.2f}, which was clamped "
        f"to {clamped_sentiment_sum:.2f} to prevent one news batch from excessively "
        f"distorting the valuation scenarios."
    )

    methodology_text = (
        "Each headline was scored using a local keyword-based sentiment model. "
        "Positive finance/event terms increased score, negative risk/event terms reduced score. "
        "Article-level sentiment weights were capped within a small range, then aggregated across "
        "the latest news set. The total score was further clamped to [-0.15, 0.15] before being "
        "used in scenario valuation."
    )

    return {
        "article_count": len(news_items),
        "positive_count": positive_count,
        "neutral_count": neutral_count,
        "negative_count": negative_count,
        "raw_sentiment_sum": raw_sentiment_sum,
        "clamped_sentiment_sum": clamped_sentiment_sum,
        "average_sentiment_weight": average_sentiment_weight,
        "sentiment_label": sentiment_label,
        "sentiment_explanation": sentiment_explanation,
        "methodology_text": methodology_text
    }




def get_options_snapshot(ticker: str) -> Dict[str, Any]:
    snap = {"expiration": None, "notes": ""}
    try:
        t = yf.Ticker(ticker)
        expirations = t.options
        if not expirations:
            snap["notes"] = "No options available."
            return snap
        exp = expirations[0]
        chain = t.option_chain(exp)
        for name, df in [("call", chain.calls), ("put", chain.puts)]:
            if not df.empty:
                df["openInterest"] = pd.to_numeric(df["openInterest"], errors="coerce").fillna(0)
                top = df.sort_values(["openInterest", "volume"], ascending=False).iloc[0]
                snap[f"top_{name}_oi_strike"] = float(top["strike"])
                snap[f"top_{name}_oi"] = int(top["openInterest"])
        snap["expiration"] = exp
        return snap
    except Exception as e:
        snap["notes"] = f"Failed: {e}"
        return snap

# ----------------------------
# 4) RAG Corpus Building
# ----------------------------
def _build_rag_corpus(ticker: str, fundamentals: Dict[str, Any], options_snap: Dict[str, Any], news_items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    corpus = []
    corpus.append({
        "text": f"Business Summary: {fundamentals.get('Business Summary')}",
        "meta": {"source": "Company Profile"}
    })
    corpus.append({
        "text": f"Derivatives Positioning: Nearest Expiration {options_snap.get('expiration')}. Top Call OI Strike at {options_snap.get('top_call_oi_strike')}. Top Put OI Strike at {options_snap.get('top_put_oi_strike')}.",
        "meta": {"source": "Options Data"}
    })
    for item in news_items:
        corpus.append({
            "text": f"LATEST NEWS: {item['title']} | TONE: {item['tone_tag']} | SOURCE: {item['publisher']} | URL: {item['link']}",
            "meta": {"source": item['publisher'], "url": item['link']}
        })
    return corpus

# ----------------------------
# 5) AI Analysis & Integrated Report
# ----------------------------
def get_ai_investment_plan(
        ticker: str,
        fundamentals: Dict[str, Any],
        current_price: float,
        model_desc: str,
        positive_price: float,
        neutral_price: float,
        negative_price: float,
        api_key: Optional[str] = None
    ) -> str:

    model = _get_gemini_model(api_key)
    if not model: return "Missing Gemini API Key. 請在查詢介面輸入您的 Gemini API Key。"

    reporting_date = datetime.now().strftime("%B %d, %Y")
    news_data = get_recent_news(ticker, max_items=10, api_key=api_key)
    opts = get_options_snapshot(ticker)
    sentiment_summary = summarize_news_sentiment(news_data)
    
    # build up a list Markdown
    news_section_md = "\n".join([
    f"* [{n['title']}]({n['link']}) - {n['publisher']} | {n['date']} | {n['tone_tag']}"
    for n in news_data
    ])

    news_scoring_md = "\n".join([
        (
            f"* {n['date']} | {n['publisher']} | {n['tone_tag']} | "
            f"raw_score={n.get('raw_score', 0)} | "
            f"weight={n.get('sentiment_weight', 0.0):.2f} | "
            f"title={n['title']}"
        )
        for n in news_data
    ])

    corpus = _build_rag_corpus(ticker, fundamentals, opts, news_data)
    
    # RAG 
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
    texts = [d["text"] for d in corpus]
    vectorizer = TfidfVectorizer(stop_words="english")
    doc_matrix = vectorizer.fit_transform(texts)
    q_vec = vectorizer.transform([f"{ticker} catalysts outlook"])
    scores = cosine_similarity(q_vec, doc_matrix).ravel()
    top_docs = [corpus[int(i)] for i in scores.argsort()[::-1][:6]]
    
    context = "\n---\n".join([f"Source: {d['meta']['source']}\n{d['text']}" for d in top_docs])

    prompt = f"""
Role: 'StockMaster'.
Ticker: {ticker} | Price: {current_price} | Profile: {model_desc}

Provide an INTEGRATED ANALYSIS REPORT. 
DO NOT use bracketed citations like [1] or [2]. 

Reporting Date: {reporting_date}

## 1. Latest Market News (Top 10)
- The reporting date for this analysis is exactly: {reporting_date}
- Do NOT infer or replace the reporting date with any article publication date.
- Only use news retrieved within the most recent 6 months.
- Present the newest articles first.
{news_section_md}

## 2. News Sentiment Scoring Methodology
Use the following methodology exactly in your explanation:
- A local keyword sentiment model was applied to each article's headline and description.
- Positive finance/event terms increase score.
- Negative risk/event terms decrease score.
- Each article receives a capped sentiment weight.
- Aggregate sentiment is computed as the sum of article sentiment weights.
- The final aggregate sentiment score is clamped to the range [-0.15, 0.15] before being used in valuation.

Sentiment summary:
- Articles analyzed: {sentiment_summary['article_count']}
- Positive articles: {sentiment_summary['positive_count']}
- Neutral articles: {sentiment_summary['neutral_count']}
- Negative articles: {sentiment_summary['negative_count']}
- Raw sentiment sum: {sentiment_summary['raw_sentiment_sum']:.2f}
- Clamped sentiment sum: {sentiment_summary['clamped_sentiment_sum']:.2f}
- Average sentiment weight: {sentiment_summary['average_sentiment_weight']:.4f}
- Overall sentiment conclusion: {sentiment_summary['sentiment_label']}

Sentiment scoring details by article:
{news_scoring_md}

You must explicitly explain:
1. how the sentiment score was calculated,
2. what the aggregate sentiment conclusion is,
3. whether the current news flow is supportive, neutral, or adverse for the stock.

## 3. NLP Sentiment & Trending Impact
- Synthesize the overall tone from the scored news articles above.
- State the final sentiment conclusion clearly (for example: moderately positive / slightly negative / neutral).
- Explain how these catalysts and current option walls ({opts.get('top_call_oi_strike', 'N/A')} call / {opts.get('top_put_oi_strike', 'N/A')} put) affect market trend.
- Explicitly explain how sentiment affects the scenario weighting.
- Distinguish between structural valuation drivers and news-driven sentiment drivers.

## 4. Options & Derivatives Analysis

## 5. Institutional Execution Plan

## 6. Simulated Price Forward Assessment
Use the following scenario prices exactly as the scenario anchors:
- Current Close: ${current_price:.2f}
- Positive Scenario (Bullish sentiment): ${positive_price:.2f}
- Neutral Scenario (Base case): ${neutral_price:.2f}
- Negative Scenario (Bearish sentiment): ${negative_price:.2f}

For this section:
- Do NOT output short-term range.
- Do NOT output long-term range.

- Use the company's sector, industry structure, market capitalization, profitability profile, and recent news sentiment to explain why the pricing framework selected the valuation model: {model_desc}.

- Your explanation should include:
  1. Why this valuation model is appropriate for the company type (for example growth, financial cycle, defensive, or commodity cycle).
  2. Which fundamental factors most influenced the scenario spread (such as revenue growth, valuation multiples, market structure, or cyclicality).
  3. How news sentiment influenced the weighting of the scenarios and why the positive, neutral, and negative prices differ.

- The explanation should read like a short institutional equity research note and remain consistent with the valuation profile: {model_desc}.

ADDITIONAL CONTEXT:
{context}
"""
    try:
        response = model.generate_content(prompt)
        return response.text or "Error: Empty response"
    except Exception as e: return f"Analysis Failed: {e}"

# ----------------------------
# 6) DOCX generation
# ----------------------------
from docx import Document
from docx.shared import Pt, RGBColor

def generate_docx_report(ticker, fundamentals, ai_content):
    doc = Document()
    doc.add_heading(f"StockMaster Research: {ticker}", 0)
    
    doc.add_heading("I. Fundamental Summary", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text, hdr_cells[1].text = "Metric", "Value"
    
    metrics = ["P/E Ratio (TTM)", "Revenue Growth (YoY)", "ROE (Return on Equity)", "Debt/Equity"]
    for m in metrics:
        row = table.add_row().cells
        row[0].text, row[1].text = m, str(fundamentals.get(m, "N/A"))

    doc.add_heading("II. Integrated Analysis & News Feed", level=1)
    for line in ai_content.split("\n"):
        if line.startswith("##"): doc.add_heading(line.replace("##", "").strip(), level=2)
        elif line.startswith("###"): doc.add_heading(line.replace("###", "").strip(), level=3)
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**"):
                    run = p.add_run(part.replace("**", ""))
                    run.bold = True
                else: p.add_run(part)

    doc.add_heading("Disclaimer", level=2)
    dis = doc.add_paragraph("This AI-generated report is for research only and does not constitute financial advice.")
    run = dis.runs[0]
    run.font.color.rgb, run.font.size = RGBColor(128, 128, 128), Pt(9)

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


from datetime import datetime
from collections import defaultdict
from typing import Any, Dict, List, Optional

import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


today = datetime.utcnow().strftime("%B %d, %Y")


today = datetime.utcnow().strftime("%B %d, %Y")

def get_ai_chat_response(ticker, question, api_key=None):
    """
    Chatbot answer pipeline with retrieval + inline citations.

    Supports:
    - static company facts: CEO, sector, industry, headquarters, business summary
    - recent/live-ish evidence: latest market snapshot, recent news, options positioning
    """
    ticker = str(ticker or "").strip().upper()
    question = str(question or "").strip()

    if not ticker:
        return "Please provide a stock ticker."

    if not question:
        return "Please enter a question."

    model = _get_gemini_model(api_key)
    if not model:
        return "Gemini API key missing. 請在查詢介面輸入您的 Gemini API Key。"

    company_name = ticker
    fundamentals = {}
    df_price = None
    df_financials = None
    news_items = []
    options_snap = {}

    # 1) fetch evidence
    try:
        df_price, fundamentals = get_stock_data(ticker)
        if fundamentals:
            company_name = fundamentals.get("Company Name", ticker) or ticker
    except Exception:
        fundamentals = {}
        company_name = ticker

    try:
        df_financials = get_financial_statements(ticker)
    except Exception:
        df_financials = None

    try:
        news_items = get_recent_news(ticker, max_items=10, api_key=api_key) or []
    except Exception:
        news_items = []

    try:
        options_snap = get_options_snapshot(ticker) or {}
    except Exception:
        options_snap = {}

    # 2) build retrieval KB
    knowledge_base = build_chat_knowledge_base(
        ticker=ticker,
        fundamentals=fundamentals,
        df_price=df_price,
        df_financials=df_financials,
        options_snap=options_snap,
        news_items=news_items
    )

    if not knowledge_base:
        return (
            f"I could not build a retrieval context for {ticker}. "
            f"Please verify the ticker symbol and try again."
        )

    # 3) retrieve relevant docs
    retrieved_docs = retrieve_chat_documents(
        user_query=question,
        ticker=ticker,
        company_name=company_name,
        knowledge_base=knowledge_base,
        top_k=5
    )

    if not retrieved_docs:
        return (
            f"I could not retrieve relevant evidence for your question about {ticker}. "
            f"Please try a more specific question."
        )

    # 4) grounded answer (retrieved KB + live Google Search fallback, scoped to this ticker)
    try:
        return generate_chat_answer_with_citations(
            model=model,
            ticker=ticker,
            company_name=company_name,
            question=question,
            retrieved_docs=retrieved_docs,
            today=today,
            use_search_grounding=True
        )
    except Exception as e:
        return f"AI Error: {e}"